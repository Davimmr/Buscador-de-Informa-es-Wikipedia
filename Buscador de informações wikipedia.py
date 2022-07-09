import wikipedia
import re
from docx import Document

name = input('Digite seu nome: ')
wikipedia.set_lang('pt')
tittle = input('Digite o Tema de sua pesquisa: \n')
while True:
    try:
        wiki = wikipedia.page(tittle)
        break
    except:
        print('Tema inválido ou não existente na Wikipedia')
        tittle = input('Digite outro Tema: \n')

text = wiki.content
text = re.sub(r'==', '', text)
text = re.sub(r'=', '', text)
text = re.sub(r'\n', '\n  ', text)
split = text.split('Veja também', 1)
text = split[0]
print(text)

document = Document()
paragraph = document.add_heading(tittle,0)
paragraph.alignment = 1
paragraph = document.add_paragraph('    ' + text)
paragraph = document.add_paragraph(name)
paragraph.alignment = 2
document.save(tittle + ".docx")
input()